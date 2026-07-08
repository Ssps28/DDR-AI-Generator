"""
report_generator.py

Generates the final DDR Report (.docx)

Author: Sujay S Pattar
"""

import os
import re

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DDRReportGenerator:

    def __init__(self, output_path="output/DDR_Report.docx"):

        self.output_path = output_path

        os.makedirs("output", exist_ok=True)

        self.doc = Document()

    def add_title(self):

        title = self.doc.add_heading(
            "Detailed Diagnostic Report (DDR)",
            level=1
        )

        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def property_summary(self, data):

        self.doc.add_heading(
            "1. Property Issue Summary",
            level=2
        )

        self.doc.add_paragraph(
            data.get(
                "property_issue_summary",
                "Not Available"
            )
        )

    def observations(self, data):

        self.doc.add_heading(
            "2. Area-wise Observations",
            level=2
        )

        observations = data.get(
            "area_wise_observations",
            []
        )

        for index, obs in enumerate(observations, start=1):

            self.doc.add_heading(
                f"{index}. {obs.get('area','Not Available')}",
                level=3
            )

            p = self.doc.add_paragraph()
            p.add_run("Observation: ").bold = True
            p.add_run(obs.get("observation", "Not Available"))

            p = self.doc.add_paragraph()
            p.add_run("Thermal Finding: ").bold = True
            p.add_run(obs.get("thermal_finding", "Not Available"))

            p = self.doc.add_paragraph()
            p.add_run("Root Cause: ").bold = True
            p.add_run(obs.get("root_cause", "Not Available"))

            p = self.doc.add_paragraph()
            p.add_run("Severity: ").bold = True
            p.add_run(obs.get("severity", "Not Available"))

            p = self.doc.add_paragraph()
            p.add_run("Reason: ").bold = True
            p.add_run(obs.get("severity_reason", "Not Available"))

            p = self.doc.add_paragraph()
            p.add_run("Recommended Action: ").bold = True
            p.add_run(obs.get("recommended_action", "Not Available"))

            self.doc.add_heading("Images", level=4)

            # Support both keys
            images = (
                obs.get("image_reference")
                or obs.get("images")
                or []
            )

            if not images:

                self.doc.add_paragraph("Image Not Available")

            else:

                for img in images:

                    found = False

                    match = re.search(r"\d+", str(img))

                    if match:

                        number = match.group()

                        for file in os.listdir("extracted_images"):

                            if f"img_{number}." in file:

                                try:

                                    self.doc.add_picture(
                                        os.path.join(
                                            "extracted_images",
                                            file
                                        ),
                                        width=Inches(2.5)
                                    )

                                    found = True
                                    break

                                except Exception:
                                    pass

                    if not found:

                        self.doc.add_paragraph(
                            f"{img} (Image Not Available)"
                        )

            self.doc.add_paragraph(
                "------------------------------------------------------------"
            )

    def additional_notes(self, data):

        self.doc.add_heading(
            "3. Additional Notes",
            level=2
        )

        self.doc.add_paragraph(
            data.get(
                "additional_notes",
                "Not Available"
            )
        )

    def missing_information(self, data):

        self.doc.add_heading(
            "4. Missing Information",
            level=2
        )

        missing = data.get(
            "missing_information",
            []
        )

        if not missing:

            self.doc.add_paragraph("None")

        else:

            for item in missing:

                self.doc.add_paragraph(
                    item,
                    style="List Bullet"
                )

    def generate(self, final_json):

        self.add_title()

        self.property_summary(final_json)

        self.observations(final_json)

        self.additional_notes(final_json)

        self.missing_information(final_json)

        self.doc.save(self.output_path)

        print("\nReport saved at:")
        print(self.output_path)